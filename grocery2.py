import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from prompts import raw_ingredients_system_prompt
import io
# import os

# Load environment variables
load_dotenv()

st.title("Shopping List Wizard")

doc_file = st.file_uploader("Upload a file")  

def read_word_file_from_upload(uploaded_file):
    """Read Word document from Streamlit uploaded file"""
    doc = Document(uploaded_file)
    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(row_data))

    return "\n".join(parts)

# Process uploaded file
if doc_file is not None:
    # Read the uploaded file
    doc_text = read_word_file_from_upload(doc_file)
    st.write("📄 **File uploaded successfully!**")
    
    # Show processing status
    st.write("🔄 **Processing your menu...**")
    
    # Initialize OpenAI client
    client = OpenAI()
    
    # Raw list of ingredients
    with st.spinner("Making raw list of ingredients..."):
        response = client.chat.completions.create(
            model="gpt-5", 
            messages=[
                {"role": "system", "content": raw_ingredients_system_prompt},
                {"role": "user", "content": doc_text}
            ],
            reasoning_effort="high"
        )
    
    answer = response.choices[0].message.content
    st.write("✅ **Initial shopping list generated!**")
    st.text_area("Initial Shopping List", answer, height=300)
    


else:
    st.write("👆 Please upload a Word document (.docx) to get started!")








   
#     # INITIAL SHOPPING LIST
#     with st.spinner("Generating initial shopping list..."):
#         response = client.chat.completions.create(
#             model="gpt-5",
#             reasoning_effort="high"   # Using gpt-4o instead of gpt-5
#             messages=[
#                 {"role": "system", "content": shopping_list_system_prompt},
#                 {"role": "user", "content": doc_text}
#             ]
#         )
    
#     initial_answer = response.choices[0].message.content
#     st.write("✅ **Initial shopping list generated!**")
#     st.text_area("Initial Shopping List", initial_answer, height=300)


#     # ERROR EVALUATION - 3 ROUNDS - check that all ingredients are present in the shopping list
#     st.write("🔍 **Running error evaluation (6 rounds)...**")
    
#     # Evaluator 1
#     with st.spinner("Evaluator 1/6 running..."):
#         response = client.chat.completions.create(
#             model="gpt-5",
# reasoning_effort="low" 
#             messages=[
#                 {"role": "system", "content": evaluator_system_prompt},
#                 {"role": "user", "content": f"Меню: {doc_text}\n Шоппинг-лист: {answer}"}
#             ]
#         )
    
#     answer1 = response.choices[0].message.content

#     # Evaluator 2
#     with st.spinner("Evaluator 2/6 running..."):
#         response = client.chat.completions.create(
#             model="gpt-5",
#     reasoning_effort="low"   # Using gpt-4o-mini instead of gpt-5-nano
#             messages=[
#                 {"role": "system", "content": evaluator_system_prompt},
#                 {"role": "user", "content": f"Меню: {doc_text}\n Шоппинг-лист: {answer}"}
#             ]
#         )
    
#     answer2 = response.choices[0].message.content

#     # Evaluator 3
#     with st.spinner("Evaluator 3/6 running..."):
#         response = client.chat.completions.create(
#             model="gpt-5",
#             reasoning_effort="high" 
#             messages=[
#                 {"role": "system", "content": evaluator_system_prompt},
#                 {"role": "user", "content": f"Меню: {doc_text}\n Шоппинг-лист: {answer}"}
#             ]
#         )
    
#     answer3 = response.choices[0].message.content

#     # Evaluator 4
#     with st.spinner("Evaluator 4/6 running..."):
#         response = client.chat.completions.create(
#             model="gpt-5",
# reasoning_effort="high"   # Using gpt-4o-mini instead of gpt-5-nano
#             messages=[
#                 {"role": "system", "content": evaluator_system_prompt},
#                 {"role": "user", "content": f"Меню: {doc_text}\n Шоппинг-лист: {answer}"}
#             ]
#         )
    
#     answer4 = response.choices[0].message.content

#     # Evaluator 5
#     with st.spinner("Evaluator 5/6 running..."):
#         response = client.chat.completions.create(
#             model="gpt-5",
# reasoning_effort="high"   # Using gpt-4o-mini instead of gpt-5-nano
#             messages=[
#                 {"role": "system", "content": evaluator_system_prompt},
#                 {"role": "user", "content": f"Меню: {doc_text}\n Шоппинг-лист: {answer}"}
#             ]
#         )
    
#     answer5 = response.choices[0].message.content

#     # Evaluator 6
#     with st.spinner("Evaluator 6/6 running..."):
#         response = client.chat.completions.create(
#             model="gpt-5",
# reasoning_effort="high"   # Using gpt-4o-mini instead of gpt-5-nano
#             messages=[
#                 {"role": "system", "content": evaluator_system_prompt},
#                 {"role": "user", "content": f"Меню: {doc_text}\n Шоппинг-лист: {answer}"}
#             ]
#         )
    
#     answer6 = response.choices[0].message.content

#     # EVALUATION SUMMARY
#     st.write("📊 **Summarizing evaluation results...**")
    
#     with st.spinner("Creating evaluation summary..."):
#         response = client.chat.completions.create(
#             model="gpt-5",
# reasoning_effort="high"   # Using gpt-4o instead of gpt-5
#             messages=[
#                 {"role": "system", "content": evaluation_summary_system_prompt},
#                 {"role": "user", "content": f"Список 1: {answer1}\nСписок 2: {answer2}\nСписок 3: {answer3}\nСписок 4: {answer4}\nСписок 5: {answer5}\nСписок 6: {answer6}"}
#             ]
#         )
    
#     eval_answer = response.choices[0].message.content
#     st.write("📋 **Evaluation Summary:**")
#     st.text_area("Summary of Errors Found", eval_answer, height=200)

#     # ERROR FIXING
#     st.write("🔧 **Starting error fixing process...**")
    
#     # Error Evaluation
#     with st.spinner("Validating errors..."):
#         response = client.chat.completions.create(
#             model="gpt-5",
# reasoning_effort="high"   # Using gpt-4o instead of gpt-5
#             messages=[
#                 {"role": "system", "content": error_evaluation_system_prompt},
#                 {"role": "user", "content": f"Меню: {doc_text}\nШоппинг-лист: {answer}\nСписок ошибок: {eval_answer}"}
#             ]
#         )
    
#     eval_answer_checked = response.choices[0].message.content
#     st.write("✅ **Validated Errors:**")
#     st.text_area("Confirmed Errors", eval_answer_checked, height=200)

#     # Error fixing
#     with st.spinner("Fixing errors in shopping list..."):
#         response = client.chat.completions.create(
#             model="gpt-5",
# reasoning_effort="high"   # Using gpt-4o instead of gpt-5
#             messages=[
#                 {"role": "system", "content": error_fixing_system_prompt},
#                 {"role": "user", "content": f"Шоппинг-лист: {answer}\nСписок ошибок: {eval_answer_checked}"}
#             ]
#         )
    
#     fixed_answer = response.choices[0].message.content
#     st.write("🔨 **Fixed Shopping List:**")
#     st.text_area("Corrected Shopping List", fixed_answer, height=300)

#     # Error fixing check
#     with st.spinner("Final verification..."):
#         response = client.chat.completions.create(
#             model="gpt-5",
# reasoning_effort="high"   # Using gpt-4o instead of gpt-5
#             messages=[
#                 {"role": "system", "content": error_fixing_check_system_prompt},
#                 {"role": "user", "content": f"Исправленный шоппинг-лист: {fixed_answer}\nСписок ошибок: {eval_answer}"}
#             ]
#         )

#     final_answer = response.choices[0].message.content
#     st.write("🎉 **Final Shopping List**")
#     st.text_area("Final Shopping List", final_answer, height=400)

#     # Add quantities
#     with st.spinner("Adding quantities..."):
#         response = client.chat.completions.create(
#             model="gpt-5",
# reasoning_effort="high"   # Using gpt-4o instead of gpt-5
#             messages=[
#                 {"role": "system", "content": quantities_system_prompt},
#                 {"role": "user", "content": f"Шоппинг-лист: {final_answer}"}
#             ]
#         )
    
#     final_answer_quantities = response.choices[0].message.content
#     st.write("🎉 **Final Shopping List with Quantities**")
#     st.text_area("Final Shopping List with Quantities", final_answer_quantities, height=400)

    
#     # Create Word document for download
#     def create_word_document(content):
#         """Create a Word document from the shopping list content"""
#         doc = Document()
        
#         # Add title
#         title = doc.add_heading('Shopping List', 0)
        
#         # Split content into lines and add as paragraphs
#         lines = content.split('\n')
#         for line in lines:
#             if line.strip():
#                 if line.startswith('I.') or line.startswith('II.') or line.startswith('III.') or line.startswith('IV.') or line.startswith('V.') or line.startswith('VI.') or line.startswith('VII.'):
#                     # Add section headers
#                     doc.add_heading(line.strip(), level=1)
#                 elif line.startswith('-') or line.startswith('•'):
#                     # Add list items
#                     doc.add_paragraph(line.strip(), style='List Bullet')
#                 else:
#                     # Add regular paragraphs
#                     doc.add_paragraph(line.strip())
        
#         return doc
    
#     # Generate Word document
#     word_doc = create_word_document(final_answer_quantities)
    
#     # Save to bytes for download
#     doc_buffer = io.BytesIO()
#     word_doc.save(doc_buffer)
#     doc_buffer.seek(0)
    
#     # Add download button for Word document
#     st.download_button(
#         "📥 Download Final Shopping List (Word Document)",
#         doc_buffer.getvalue(),
#         file_name="shopping_list.docx",
#         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#     )