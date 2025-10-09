import streamlit as st
import io
from openai import OpenAI
from docx import Document
from prompts import *

def read_word_file_from_upload(uploaded_file):
    """Read Word document from Streamlit uploaded file"""
    doc = Document(uploaded_file)
    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(row_data))

    return "\n".join(parts)

def create_word_document(content):
    """Create a Word document from the shopping list content"""
    doc = Document()
    
    # Add title
    doc.add_heading('Shopping List', 0)
    
    # Split content into lines and add as paragraphs
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if line:
            # Check for Roman numerals (I., II., III., etc.) or numbered sections
            if (line.startswith(('I.', 'II.', 'III.', 'IV.', 'V.', 'VI.', 'VII.', 'VIII.', 'IX.', 'X.')) or
                line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.'))):
                # Add section headers
                doc.add_heading(line, level=1)
            # Check for various bullet point formats
            elif line.startswith(('-', '•', '*', '◦', '▪', '▫')):
                # Add list items
                doc.add_paragraph(line, style='List Bullet')
            # Check for numbered list items
            elif line.startswith(('1)', '2)', '3)', '4)', '5)', '6)', '7)', '8)', '9)')):
                doc.add_paragraph(line, style='List Number')
            else:
                # Add regular paragraphs
                doc.add_paragraph(line)
    
    return doc

def run_raw_ingredients_round(round_num, total_rounds):
    """Run a single raw ingredients round"""
    with st.spinner(f"Raw list of ingredients {round_num}/{total_rounds} running..."):
        response = client.chat.completions.create(
            model="gpt-5",
            reasoning_effort=reasoning_effort,
            messages=[
                {"role": "system", "content": raw_ingredients_system_prompt},
                {"role": "user", "content": doc_text}
            ]
        )
    return response.choices[0].message.content

def run_evaluator(round_num, total_rounds):
    """Run a single evaluator round"""
    with st.spinner(f"Evaluator {round_num}/{total_rounds} running..."):
        response = client.chat.completions.create(
            model="gpt-5",
            reasoning_effort=reasoning_effort,
            messages=[
                {"role": "system", "content": evaluator_system_prompt},
                {"role": "user", "content": f"1) Меню: {doc_text}\n 2) Сгруппированный список ингредиентов: {grouped_answer}\n 3) Шоппинг-лист: {counts_answer}"}
            ]
        )
    return response.choices[0].message.content

def save_word_doc_to_buffer(word_doc):
    doc_buffer = io.BytesIO()
    word_doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer


# Load api key variables (stored in Streamlit secrets)
api_key = st.secrets["OPENAI_API_KEY"]
reasoning_effort = 'low'

st.title("Shopping List Wizard")

doc_file = st.file_uploader("Upload a file")  

# Process uploaded file
if doc_file is not None:
    # Read the uploaded file
    doc_text = read_word_file_from_upload(doc_file)
    
    st.write("📄 **File uploaded successfully!**")

    # Show processing status
    st.write("🔄 **Processing your menu...**")
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)
    
    # Raw list of ingredients - 3 rounds
    raw_answers = []
    for i in range(1, 4):
        answer = run_raw_ingredients_round(i, 3)
        raw_answers.append(answer)
    
    # Unpack answers for backward compatibility
    raw_answer1, raw_answer2, raw_answer3 = raw_answers

    # Final raw list of ingredients
    with st.spinner("Final raw list of ingredients running..."):
        response = client.chat.completions.create(
            model="gpt-5",
            reasoning_effort=reasoning_effort, 
            messages=[
                {"role": "system", "content": clean_raw_ingredients_system_prompt},
                {"role": "user", "content": f"Cписок 1: {raw_answer1}\nCписок 2: {raw_answer2}\nCписок 3: {raw_answer3}"}
            ]
        )
    final_raw_answer = response.choices[0].message.content
    st.text_area("Final raw list of ingredients", final_raw_answer, height=400)
    
    # Normalized list of ingredients
    with st.spinner("Normalized list of ingredients running..."):
        response = client.chat.completions.create(
            model="gpt-5",
                reasoning_effort=reasoning_effort, 
            messages=[
                {"role": "system", "content": normalized_ingredients_system_prompt},
                {"role": "user", "content": final_raw_answer}
            ]
        )
    normalized_answer = response.choices[0].message.content

    # Grouped list of ingredients
    with st.spinner("Grouped list of ingredients running..."):
        response = client.chat.completions.create(
            model="gpt-5",
            reasoning_effort=reasoning_effort, 
            messages=[
                {"role": "system", "content": grouped_ingredients_system_prompt},
                {"role": "user", "content": normalized_answer}
            ]
        )
    grouped_answer = response.choices[0].message.content
    st.text_area("Grouped list of ingredients", grouped_answer, height=400)

      
    # Counts of ingredients
    with st.spinner("Counts of ingredients running..."):
        response = client.chat.completions.create(
            model="gpt-5",
            reasoning_effort=reasoning_effort, 
            messages=[
                {"role": "system", "content": counts_system_prompt},
                {"role": "user", "content": f"Menu: {doc_text}\nGrouped list of ingredients: {grouped_answer}"}
            ]
        )
    counts_answer = response.choices[0].message.content

    st.write("✅ **Counts of ingredients generated!**")
    st.text_area("Counts of ingredients", counts_answer, height=400)

    
    # ERROR EVALUATION - 6 ROUNDS
    st.write("🔍 **Running error evaluation (6 rounds)...**")
        
    # Run all 6 evaluator rounds
    answers = []
    for i in range(1, 7):
        answer = run_evaluator(i, 6)
        answers.append(answer)
    
    # Unpack answers for backward compatibility
    answer1, answer2, answer3, answer4, answer5, answer6 = answers

    # Evaluation summary
    with st.spinner("Evaluation summary running..."):
        response = client.chat.completions.create(
            model="gpt-5",
            reasoning_effort=reasoning_effort,
            messages=[{"role": "system", "content": evaluation_summary_system_prompt}, {"role": "user", "content": f"Результат 1: {answer1}\nРезультат 2: {answer2}\nРезультат 3: {answer3}\nРезультат 4: {answer4}\nРезультат 5: {answer5}\nРезультат 6: {answer6}"}]
        )
    evaluation_summary = response.choices[0].message.content
    st.text_area("Evaluation summary", evaluation_summary, height=400)

    # Generate Word document
    word_doc = create_word_document(counts_answer)
    doc_buffer = save_word_doc_to_buffer(word_doc)
    
    # Add download button for Word document
    st.download_button(
        "📥 Download Final Shopping List (Word Document)",
        doc_buffer.getvalue(),
        file_name=f"Шоппинг.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

else:
    st.write("👆 Please upload a Word document (.docx) to get started!")