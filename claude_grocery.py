import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
import pandas as pd
import json
from typing import List, Dict
import concurrent.futures

load_dotenv()

st.title("🛒 Maximum Recall Grocery List Generator")

# ============================================================================
# CONFIGURATION
# ============================================================================

EXTRACTION_SCHEMA = {
    "type": "json_schema",
    "json_schema": {
        "name": "ingredient_extraction",
        "strict": True,
        "schema": {
            "type": "object",
            "properties": {
                "ingredients": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "quantity": {"type": "number"},
                            "unit": {"type": "string"},
                            "day": {"type": "integer"},
                            "meal": {"type": "string"},
                            "dish": {"type": "string"}
                        },
                        "required": ["name", "quantity", "unit", "day", "meal", "dish"],
                        "additionalProperties": False
                    }
                }
            },
            "required": ["ingredients"],
            "additionalProperties": False
        }
    }
}

EXTRACTION_PROMPT = """Извлеки ВСЕ ингредиенты из меню в JSON формате.

КРИТИЧЕСКИ ВАЖНО - НЕ ПРОПУСТИ НИ ОДНОГО ИНГРЕДИЕНТА:
1. Включи КАЖДЫЙ ингредиент, даже если "по желанию", "по вкусу", "украсить"
2. Включи ВСЕ специи, масла, соусы, зелень - даже если кажутся необязательными
3. Используй ТОЧНЫЕ количества из рецепта (уже на все порции - НЕ умножай!)
4. Для "десерт на 150 ккал" или "снеки на 200 ккал" - включи дословно как есть
5. НЕ включай: только воду для варки (но ВКЛЮЧИ покупной бульон)
6. Любое молоко → "растительное молоко"

Для каждого ингредиента:
- name: нормализованное название (помидоры/томаты → "томаты", твёрдый сыр → "сыр")
- quantity: точное число из рецепта
- unit: точная единица (г, мл, шт, стакан, ст.л., ч.л.)
- day: номер дня (1-7)
- meal: завтрак/обед/перекус/ужин
- dish: краткое название блюда

ПРОВЕРЬ ДВАЖДЫ: Если блюдо готовится на 3 порции и используется 3 дня подряд - 
укажи ингредиенты ОДИН раз для дня готовки, НЕ умножай на количество дней!"""

# ============================================================================
# CATEGORIES
# ============================================================================

CATEGORIES = {
    "МЯСО, ПТИЦА, РЫБА, МОРЕПРОДУКТЫ": [
        "говяд", "курин", "треска", "креветк", "фарш", "филе", "рыб", "морепродукт"
    ],
    "ЯЙЦА И МОЛОЧНЫЕ ПРОДУКТЫ": [
        "яйц", "творог", "йогурт", "сыр", "молоко", "сливк", "пармезан", "кефир"
    ],
    "КРУПЫ, ЗЛАКИ, БОБОВЫЕ, МУКА": [
        "гречк", "рис", "перловк", "чечевиц", "булгур", "мука", "кукурузн", "полб", 
        "хлеб", "лаваш", "хлебц", "овсян", "бобов"
    ],
    "ОВОЩИ, ГРИБЫ, ЗЕЛЕНЬ": [
        "томат", "помидор", "огурец", "перец", "лук", "чеснок", "морков",
        "баклажан", "кабачок", "шпинат", "грибы", "шампиньон", "картофель",
        "руккола", "зелень", "петрушк", "укроп", "кинза", "базилик", "тыкв",
        "салат", "капуст", "свекл", "вешенк"
    ],
    "ФРУКТЫ И ЯГОДЫ": [
        "яблок", "клубник", "ягод", "фрукт", "лимон", "лайм", "авокадо", "груш", "банан"
    ],
    "ОРЕХИ И СЕМЕНА": [
        "орех", "фундук", "грецк", "миндал", "семен", "кунжут", "чиа", "мак"
    ],
    "ДЕСЕРТЫ И СНЕКИ": [
        "десерт", "снек", "батончик", "икра", "паста ореховая", "протеин", "вафл"
    ],
    "СОУСЫ, МАСЛА, СПЕЦИИ": [
        "масло", "уксус", "соус", "ткемали", "горчиц", "паста томатн",
        "специи", "соль", "перец", "паприк", "кориандр", "хмели",
        "куркума", "тимьян", "орегано", "корица", "разрыхлитель", "сахар",
        "гхи", "оливков"
    ]
}

def classify_category(name: str) -> str:
    name_lower = name.lower()
    for category, keywords in CATEGORIES.items():
        if any(kw in name_lower for kw in keywords):
            return category
    return "ПРОЧИЕ ИНГРЕДИЕНТЫ"

# ============================================================================
# EXTRACTION FUNCTIONS
# ============================================================================

def read_word_file(uploaded_file):
    doc = Document(uploaded_file)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(row_data))
    return "\n".join(parts)


def extract_with_gpt4o(client: OpenAI, menu_text: str, attempt: int = 1) -> List[Dict]:
    """Extract with GPT-4o"""
    response = client.chat.completions.create(
        model="gpt-4o",
        temperature=0.3 if attempt == 1 else 0.7,  # Vary temperature
        messages=[
            {"role": "system", "content": EXTRACTION_PROMPT},
            {"role": "user", "content": menu_text}
        ],
        response_format=EXTRACTION_SCHEMA
    )
    result = json.loads(response.choices[0].message.content)
    return result["ingredients"]


def extract_with_o1(client: OpenAI, menu_text: str) -> List[Dict]:
    """Extract with o1-mini for careful reasoning"""
    response = client.chat.completions.create(
        model="o1-mini",
        messages=[
            {"role": "user", "content": f"{EXTRACTION_PROMPT}\n\nМЕНЮ:\n{menu_text}"}
        ]
    )
    
    # o1 doesn't support structured outputs, parse JSON from text
    content = response.choices[0].message.content
    # Find JSON in response
    start = content.find('[')
    end = content.rfind(']') + 1
    if start != -1 and end > start:
        ingredients_array = json.loads(content[start:end])
        return ingredients_array
    return []


def extract_by_day(client: OpenAI, menu_text: str, day: int) -> List[Dict]:
    """Extract ingredients for a specific day - reduces context"""
    lines = menu_text.split('\n')
    
    # Find day section
    day_start = -1
    day_end = len(lines)
    
    for i, line in enumerate(lines):
        if f'День {day}' in line or f'Day {day}' in line:
            day_start = i
        elif day_start != -1 and (f'День {day+1}' in line or f'Day {day+1}' in line):
            day_end = i
            break
    
    if day_start == -1:
        return []
    
    day_menu = '\n'.join(lines[day_start:day_end])
    
    response = client.chat.completions.create(
        model="gpt-4o",
        temperature=0.1,
        messages=[
            {"role": "system", "content": EXTRACTION_PROMPT + f"\n\nЭКСТРАГИРУЙ ТОЛЬКО ДЛЯ ДНЯ {day}!"},
            {"role": "user", "content": day_menu}
        ],
        response_format=EXTRACTION_SCHEMA
    )
    result = json.loads(response.choices[0].message.content)
    return result["ingredients"]


# ============================================================================
# VERIFICATION FUNCTIONS
# ============================================================================

def verify_completeness(client: OpenAI, menu_text: str, extracted_names: List[str]) -> List[str]:
    """Ask model to list any missing ingredients"""
    response = client.chat.completions.create(
        model="gpt-4o",
        temperature=0,
        messages=[
            {"role": "system", "content": f"""Меню:\n{menu_text[:4000]}

Извлеченные ингредиенты: {', '.join(extracted_names)}

Верни JSON ТОЛЬКО с ингредиентами, которые есть в меню но ОТСУТСТВУЮТ в извлеченных.
НЕ включай: воду для варки, приготовленные блюда.
ВКЛЮЧИ: все специи, масла, соусы, зелень.

Формат: {{"missing": ["ingredient1", "ingredient2"]}}"""},
            {"role": "user", "content": "Какие ингредиенты пропущены?"}
        ],
        response_format={"type": "json_object"}
    )
    result = json.loads(response.choices[0].message.content)
    return result.get("missing", [])


def cross_check_with_claude(menu_text: str, extracted_names: List[str]) -> List[str]:
    """Cross-check with Claude for different perspective"""
    try:
        import anthropic
        claude = anthropic.Anthropic()
        
        response = claude.messages.create(
            model="claude-sonnet-4",
            max_tokens=2000,
            temperature=0,
            messages=[{
                "role": "user",
                "content": f"""Меню:\n{menu_text[:5000]}

Извлечённые ингредиенты: {', '.join(extracted_names)}

Верни JSON с ингредиентами из меню, которых НЕТ в извлечённом списке.
Формат: {{"missing": [...]}}"""
            }]
        )
        
        result = json.loads(response.content[0].text)
        return result.get("missing", [])
    except Exception as e:
        st.warning(f"Claude недоступен: {e}")
        return []


def verify_quantities(client: OpenAI, menu_text: str, df: pd.DataFrame) -> List[Dict]:
    """Spot-check quantities for critical items"""
    critical = df[df['category'].str.contains('МЯСО|РЫБА|КРУПЫ|МОЛОЧН')].head(15)
    
    items_text = "\n".join([
        f"{row['name']}: {row['quantity']}{row['unit']}" 
        for _, row in critical.iterrows()
    ])
    
    response = client.chat.completions.create(
        model="gpt-4o",
        temperature=0,
        messages=[
            {"role": "system", "content": f"""Меню:\n{menu_text[:4000]}

Критические ингредиенты и количества:\n{items_text}

Верни JSON с ошибками количества (только если отклонение >30%).
Формат: {{"errors": [{{"name": "x", "extracted": "100г", "should_be": "~300г"}}]}}"""},
            {"role": "user", "content": "Есть ли грубые ошибки количества?"}
        ],
        response_format={"type": "json_object"}
    )
    result = json.loads(response.choices[0].message.content)
    return result.get("errors", [])


# ============================================================================
# AGGREGATION
# ============================================================================

def aggregate_ingredients(ingredients: List[Dict]) -> pd.DataFrame:
    df = pd.DataFrame(ingredients)
    
    # Clean data types to ensure they're hashable for drop_duplicates
    def clean_field(field):
        if pd.isna(field) or field is None:
            return ""
        if isinstance(field, (list, dict)):
            return str(field)
        return str(field)
    
    # Clean all fields used in duplicate detection
    df['name'] = df['name'].apply(clean_field)
    df['day'] = df['day'].apply(clean_field)
    df['meal'] = df['meal'].apply(clean_field)
    df['dish'] = df['dish'].apply(clean_field)
    
    # Remove duplicates within same extraction
    df = df.drop_duplicates(subset=['name', 'day', 'meal', 'dish'])
    
    # Clean and convert quantity to numeric
    def clean_quantity(qty):
        if pd.isna(qty) or qty is None:
            return 0
        if isinstance(qty, (int, float)):
            return float(qty)
        if isinstance(qty, str):
            # Try to extract number from string
            import re
            numbers = re.findall(r'\d+\.?\d*', str(qty))
            if numbers:
                return float(numbers[0])
        return 0
    
    df['quantity'] = df['quantity'].apply(clean_quantity)
    
    # Filter out rows with zero quantity
    df = df[df['quantity'] > 0]
    
    if len(df) == 0:
        return pd.DataFrame(columns=['name', 'quantity', 'unit', 'day', 'meal', 'dish', 'category'])
    
    # Convert day back to numeric for proper sorting
    def parse_day(day_str):
        try:
            return int(day_str) if day_str.isdigit() else 0
        except:
            return 0
    
    df['day_numeric'] = df['day'].apply(parse_day)
    
    aggregated = df.groupby(['name', 'unit'], as_index=False).agg({
        'quantity': 'sum',
        'day_numeric': lambda x: sorted(set(x)),
        'meal': lambda x: list(set(x)),
        'dish': lambda x: list(set(x))
    })
    
    # Rename back to 'day' for consistency
    aggregated = aggregated.rename(columns={'day_numeric': 'day'})
    
    aggregated['category'] = aggregated['name'].apply(classify_category)
    return aggregated


def add_piece_counts(df: pd.DataFrame) -> pd.DataFrame:
    PRODUCE_WEIGHTS = {
        "томаты": 100, "помидор": 100, "огурец": 150, "перец": 150,
        "лук": 100, "морковь": 100, "картофель": 150, "яблок": 180,
        "авокадо": 200, "лимон": 100, "баклажан": 300, "кабачок": 250,
        "яйц": 50
    }
    
    def calc_pieces(row):
        if row['unit'] in ['г', 'гр', 'шт']:
            for produce, weight in PRODUCE_WEIGHTS.items():
                if produce in row['name'].lower():
                    if row['unit'] == 'шт':
                        return f"{int(row['quantity'])} шт"
                    pieces = int(row['quantity'] / weight + 0.9)
                    return f"{row['quantity']}{row['unit']} (~{pieces} шт)"
        return f"{row['quantity']}{row['unit']}"
    
    df['quantity_display'] = df.apply(calc_pieces, axis=1)
    return df


def format_grocery_list(df: pd.DataFrame) -> str:
    output = []
    categories_order = [
        "МЯСО, ПТИЦА, РЫБА, МОРЕПРОДУКТЫ",
        "ЯЙЦА И МОЛОЧНЫЕ ПРОДУКТЫ",
        "КРУПЫ, ЗЛАКИ, БОБОВЫЕ, МУКА",
        "ОВОЩИ, ГРИБЫ, ЗЕЛЕНЬ",
        "ФРУКТЫ И ЯГОДЫ",
        "ОРЕХИ И СЕМЕНА",
        "ДЕСЕРТЫ И СНЕКИ",
        "СОУСЫ, МАСЛА, СПЕЦИИ",
        "ПРОЧИЕ ИНГРЕДИЕНТЫ"
    ]
    
    for i, category in enumerate(categories_order, 1):
        items = df[df['category'] == category]
        if len(items) == 0:
            continue
        output.append(f"\n**{i}. {category}**\n")
        for _, row in items.iterrows():
            days_str = ", ".join([f"Д{d}" for d in row['day']])
            meals_str = ", ".join(set(row['meal']))
            dishes_str = "; ".join(list(set(row['dish']))[:3])
            output.append(
                f"- {row['name'].capitalize()} {row['quantity_display']} "
                f"({days_str} {meals_str}: {dishes_str})"
            )
    return "\n".join(output)


# ============================================================================
# MAIN APP
# ============================================================================

doc_file = st.file_uploader("📁 Загрузите меню (.docx)", type=['docx'])

if doc_file:
    menu_text = read_word_file(doc_file)
    st.success("✅ Файл загружен!")
    
    with st.expander("📄 Просмотр меню"):
        st.text(menu_text[:2000] + "..." if len(menu_text) > 2000 else menu_text)
    
    if st.button("🚀 СОЗДАТЬ СПИСОК (Максимальная точность)"):
        client = OpenAI()
        all_ingredients = []
        
        # ===== PHASE 1: MULTI-MODEL ENSEMBLE EXTRACTION =====
        st.markdown("## 🔍 Фаза 1: Извлечение (ансамбль моделей)")
        
        with st.spinner("Извлечение GPT-4o (3 параллельных попытки с разной температурой)..."):
            with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
                futures = [
                    executor.submit(extract_with_gpt4o, client, menu_text, i)
                    for i in range(1, 4)
                ]
                gpt4o_results = [f.result() for f in futures]
            
            for result in gpt4o_results:
                all_ingredients.extend(result)
            st.success(f"✅ GPT-4o: {sum(len(r) for r in gpt4o_results)} записей")
        
        with st.spinner("Извлечение o1-mini (глубокий анализ)..."):
            try:
                o1_result = extract_with_o1(client, menu_text)
                all_ingredients.extend(o1_result)
                st.success(f"✅ o1-mini: {len(o1_result)} записей")
            except Exception as e:
                st.warning(f"o1-mini пропущен: {e}")
        
        with st.spinner("Извлечение по дням (7 параллельных вызовов)..."):
            with concurrent.futures.ThreadPoolExecutor(max_workers=7) as executor:
                futures = [
                    executor.submit(extract_by_day, client, menu_text, day)
                    for day in range(1, 8)
                ]
                day_results = [f.result() for f in futures]
            
            for result in day_results:
                all_ingredients.extend(result)
            st.success(f"✅ По дням: {sum(len(r) for r in day_results)} записей")
        
        st.info(f"📊 Всего извлечено записей: {len(all_ingredients)}")
        
        # ===== PHASE 2: AGGREGATION =====
        st.markdown("## 📊 Фаза 2: Агрегация")
        
        with st.spinner("Программное суммирование количеств..."):
            df = aggregate_ingredients(all_ingredients)
            df = add_piece_counts(df)
            st.success(f"✅ Уникальных ингредиентов: {len(df)}")
        
        grocery_list = format_grocery_list(df)
        extracted_names = df['name'].unique().tolist()
        
        # ===== PHASE 3: MULTI-LAYER VERIFICATION =====
        st.markdown("## ✅ Фаза 3: Многослойная верификация")
        
        all_missing = []
        
        # Layer 1: Full menu check
        with st.spinner("Слой 1: Проверка полноты (GPT-4o)..."):
            missing_1 = verify_completeness(client, menu_text, extracted_names)
            all_missing.extend(missing_1)
            st.info(f"Найдено потенциально пропущенных: {len(missing_1)}")
        
        # Layer 2: Day-by-day verification
        with st.spinner("Слой 2: Проверка по дням (7 параллельных)..."):
            with concurrent.futures.ThreadPoolExecutor(max_workers=7) as executor:
                futures = []
                for day in range(1, 8):
                    day_ingredients = df[df['day'].apply(lambda x: day in x)]['name'].tolist()
                    futures.append(
                        executor.submit(verify_completeness, client, menu_text, day_ingredients)
                    )
                day_missing = [item for f in futures for item in f.result()]
            
            all_missing.extend(day_missing)
            st.info(f"Найдено по дням: {len(day_missing)}")
        
        # Layer 3: Cross-model check with Claude
        with st.spinner("Слой 3: Кросс-проверка (Claude)..."):
            claude_missing = cross_check_with_claude(menu_text, extracted_names)
            all_missing.extend(claude_missing)
            st.info(f"Claude нашёл: {len(claude_missing)}")
        
        # Layer 4: Quantity verification
        with st.spinner("Слой 4: Проверка количеств..."):
            qty_errors = verify_quantities(client, menu_text, df)
            if qty_errors:
                st.warning(f"⚠️ Найдено ошибок количества: {len(qty_errors)}")
                st.json(qty_errors)
        
        # Deduplicate and clean missing list
        all_missing = list(set([m.lower().strip() for m in all_missing if m]))
        
        # ===== PHASE 4: AUTO-FIX =====
        if all_missing:
            st.warning(f"⚠️ Потенциально пропущено: {len(all_missing)} ингредиентов")
            st.write(all_missing)
            
            with st.spinner("🔧 Автоматическое добавление пропущенных..."):
                # Re-extract just the missing ones
                missing_prompt = f"""Из этого меню извлеки ТОЛЬКО следующие ингредиенты: {', '.join(all_missing)}

Меню:\n{menu_text}

{EXTRACTION_PROMPT}"""
                
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": missing_prompt},
                        {"role": "user", "content": "Извлеки эти пропущенные ингредиенты"}
                    ],
                    response_format=EXTRACTION_SCHEMA
                )
                
                fixed = json.loads(response.choices[0].message.content)
                all_ingredients.extend(fixed['ingredients'])
                
                # Re-aggregate
                df = aggregate_ingredients(all_ingredients)
                df = add_piece_counts(df)
                grocery_list = format_grocery_list(df)
                
                st.success(f"✅ Добавлено {len(fixed['ingredients'])} ингредиентов. Итого: {len(df)}")
        else:
            st.success("🎉 Все проверки пройдены - пропусков не обнаружено!")
        
        # ===== FINAL OUTPUT =====
        st.markdown("---")
        st.markdown("## 🛒 ФИНАЛЬНЫЙ СПИСОК ПОКУПОК")
        st.markdown(grocery_list)
        
        st.download_button(
            label="📥 Скачать список",
            data=grocery_list,
            file_name="shopping_list.txt",
            mime="text/plain"
        )
        
        # Stats
        st.markdown("### 📈 Статистика")
        col1, col2, col3 = st.columns(3)
        col1.metric("Всего ингредиентов", len(df))
        col2.metric("Записей извлечено", len(all_ingredients))
        col3.metric("Добавлено при проверке", len(all_missing))